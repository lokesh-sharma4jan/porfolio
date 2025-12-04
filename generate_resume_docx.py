from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._element.append(hyperlink)
    return hyperlink

def create_resume_docx():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    
    # Header - Name
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('LOKESH BABU SHARMA')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 54, 93)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Senior Java Developer | Cloud Architect | Microservices Expert')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 82, 130)
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Email: luckysharma824@gmail.com | Phone: +91-9997105792\n')
    run.font.size = Pt(10)
    run = contact.add_run('LinkedIn: linkedin.com/in/the-lokesh-babu | GitHub: github.com/luckysharma824 | Location: Gurgaon, Haryana, India')
    run.font.size = Pt(10)
    
    # Professional Summary
    doc.add_paragraph()
    heading = doc.add_heading('PROFESSIONAL SUMMARY', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    summary = doc.add_paragraph(
        'Results-driven Senior Technology Consultant with 6.5+ years of progressive experience architecting and delivering '
        'enterprise-scale Java applications for Fortune 500 clients. Expert in cloud-native microservices development, AWS serverless '
        'architecture, and application modernization initiatives. Proven track record of delivering $500K+ in cost savings, processing '
        '1M+ daily transactions, and achieving 80% performance improvements through strategic optimization. Specialized in transforming '
        'legacy monolithic systems into scalable cloud solutions while leading cross-functional global teams. AWS Certified Cloud '
        'Practitioner with demonstrated expertise in Spring Boot, Hibernate, Event-Driven Architecture, and DevOps practices.'
    )
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Core Competencies
    heading = doc.add_heading('CORE COMPETENCIES', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    competencies = [
        ('Programming Languages:', 'Java (Core & Advanced), J2EE, JSP, Servlets, JDBC, Multithreading'),
        ('Frameworks & Libraries:', 'Spring Boot, Spring MVC, Spring Security, Spring Cloud, Hibernate, JPA, gRPC'),
        ('Microservices & Architecture:', 'RESTful APIs, Event-Driven Design, Domain-Driven Design, SOLID Principles'),
        ('Cloud Platforms:', 'AWS (Lambda, EC2, S3, SQS, SNS, DynamoDB, RDS, CloudWatch, Step Functions)'),
        ('Databases:', 'MySQL, PostgreSQL, Oracle, MongoDB, DynamoDB, Cassandra, Redis, RocksDB'),
        ('DevOps & Tools:', 'Docker, Kubernetes, Jenkins, Git, Maven, Gradle, Terraform, CloudFormation'),
        ('Messaging & Streaming:', 'Apache Kafka, IBM MQ, AWS SQS/SNS, Event Streaming'),
        ('Testing & Quality:', 'JUnit, Mockito, Test-Driven Development (TDD), SonarQube, Code Reviews'),
        ('Monitoring & Logging:', 'ELK Stack (Elasticsearch, Logstash, Kibana), CloudWatch, APM'),
        ('Methodologies:', 'Agile/Scrum, CI/CD Pipelines, DevOps, Code Quality Management, Technical Mentorship')
    ]
    
    for title, skills in competencies:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title} ')
        run.bold = True
        p.add_run(skills)
    
    # Professional Experience
    heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    # Ernst & Young
    p = doc.add_paragraph()
    run = p.add_run('Senior Technology Consultant (Java Developer)')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 82, 130)
    p.add_run(' | Mar 2022 - Present')
    
    p = doc.add_paragraph()
    run = p.add_run('ERNST & YOUNG LLP')
    run.font.bold = True
    
    ey_bullets = [
        'Spearheaded enterprise application modernization initiatives for Fortune 500 clients, transforming legacy monolithic systems into cloud-native microservices architectures, migrating 500GB+ data and handling 25K+ concurrent users with zero downtime during deployment',
        'Architected and deployed serverless solutions using AWS Lambda, Step Functions, DynamoDB, and SQS/SNS, processing 500K+ events daily and reducing infrastructure costs by $200K annually (40%) while improving scalability and reducing latency from 2s to 200ms',
        'Led technical implementation within cross-functional agile teams of 15+ developers across 3 global locations, driving best practices in code review, CI/CD pipeline optimization, and DevOps workflows, delivering 12+ production releases with 98% on-time delivery rate',
        'Engineered 15+ mission-critical microservices serving as foundational components for 20+ dependent applications, processing 1M+ daily transactions and achieving 80% response time improvement (from 1.2s to 240ms) through strategic caching, query optimization, and distributed computing patterns',
        'Drove code quality transformation by refactoring 20+ legacy services (150K+ lines of code), reducing codebase complexity by 35%, improving maintainability scores from 40% to 85% through adherence to SOLID principles, and increasing test coverage from 45% to 92%',
        'Resolved 75+ critical production defects affecting 100K+ users through systematic root cause analysis and preventive measures, reducing application crashes by 70% (from 200 to 60 monthly incidents), improving overall system stability to 99.95% uptime, and reducing MTTR from 4 hours to 45 minutes'
    ]
    
    for bullet in ey_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    run = p.add_run('Key Technologies: ')
    run.italic = True
    run = p.add_run('Java 8/11/17, Spring Boot, AWS Lambda, DynamoDB, Terraform, CloudFormation, Oracle, PostgreSQL, Redis, Kafka, Docker, Jenkins, ELK Stack')
    run.italic = True
    
    # Smartbox
    p = doc.add_paragraph()
    run = p.add_run('Software Engineer (Java Developer)')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 82, 130)
    p.add_run(' | May 2021 - Mar 2022')
    
    p = doc.add_paragraph()
    run = p.add_run('SMARTBOX ECOMMERCE SOLUTIONS PVT. LTD.')
    run.font.bold = True
    
    smartbox_bullets = [
        'Developed core backend infrastructure for IoT-enabled smart locker logistics platform serving 500+ retail locations across 15 cities, handling 15K+ daily locker operations with 99.8% uptime SLA compliance and processing 500K+ transactions monthly',
        'Designed and implemented QR code authentication and multi-channel SMS engagement systems, processing 30K+ monthly messages, increasing customer adoption rates by 45% and improving customer satisfaction scores from 3.2 to 4.5/5 through seamless user experience',
        'Established comprehensive monitoring framework using ELK Stack (Elasticsearch, Logstash, Kibana) processing 2GB+ daily logs, enabling real-time analytics and proactive incident detection across 100+ endpoints, reducing MTTR by 60% (from 2.5 hours to 1 hour)',
        'Engineered automated audit service from ground up to track messaging and email costs across platform, monitoring $15K+ monthly communication spend, providing financial visibility that identified $60K annual cost optimization opportunities (35%)',
        'Championed code quality improvements through systematic refactoring initiatives across 80K+ lines of code, reducing technical debt by 30%, eliminating 200+ code smells, while enhancing test coverage from 55% to 85% and reducing build time by 40%'
    ]
    
    for bullet in smartbox_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    run = p.add_run('Key Technologies: ')
    run.italic = True
    run = p.add_run('Java 8, Spring Boot, MySQL, AWS S3, Redis, ELK Stack, REST APIs, Microservices, JUnit')
    run.italic = True
    
    # Oodles
    p = doc.add_paragraph()
    run = p.add_run('Senior Associate Consultant (Java Developer)')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 82, 130)
    p.add_run(' | Apr 2019 - May 2021')
    
    p = doc.add_paragraph()
    run = p.add_run('OODLES TECHNOLOGIES PVT. LTD.')
    run.font.bold = True
    
    oodles_bullets = [
        'Delivered technical solutions across diverse verticals for 8+ enterprise clients including blockchain-based identity systems, IoT platforms, social networking applications (15K+ active users), and e-commerce marketplaces processing $500K+ monthly GMV',
        'Architected microservices-based solutions serving 25K+ daily users with RESTful API design, implementing OAuth 2.0/JWT authentication and role-based access control for enterprise security compliance, achieving zero security breaches and maintaining 100% SOC 2 compliance',
        'Pioneered WhatsApp Business API integration for automated customer engagement, processing 75K+ monthly messages across 5 client projects, with 98.5% delivery success rate and reducing customer response time from 4 hours to 15 minutes, improving customer engagement by 65%',
        'Successfully delivered 50+ feature enhancements and resolved 80+ production defects across 6 concurrent projects, maintaining 95% sprint commitment achievement and zero missed SLA deadlines over 24 months',
        'Elevated engineering standards by introducing code review protocols (500+ reviews conducted), unit testing frameworks, and design patterns, improving overall code quality metrics by 50%, reducing defect density from 8 to 3 per KLOC, and improving code coverage from 40% to 90%'
    ]
    
    for bullet in oodles_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph()
    run = p.add_run('Key Technologies: ')
    run.italic = True
    run = p.add_run('Java 8, Spring Boot, PostgreSQL, MySQL, MongoDB, Redis, AWS S3, OAuth 2.0, JWT, WebRTC, WhatsApp Business API')
    run.italic = True
    
    # Key Projects
    doc.add_page_break()
    heading = doc.add_heading('KEY PROJECTS', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    projects = [
        {
            'title': 'CUSTOMER OUTPUT - FIDELITY | Ernst & Young LLP',
            'desc': 'Architected enterprise-grade automated communication platform delivering multi-channel customer notifications (email, PDF, postal). Designed serverless AWS infrastructure processing 100K+ daily transactions with 80% latency reduction through event-driven architecture.',
            'tech': 'Java, AWS Lambda, Step Functions, CloudFormation, Terraform, Oracle, DynamoDB, SQS/SNS'
        },
        {
            'title': 'SMARTBOX DIGITAL LOCKERS | Smartbox Ecommerce Solutions',
            'desc': 'Engineered IoT-enabled digital locker management system supporting 500+ retail locations. Implemented real-time device monitoring, QR-based authentication, and SMS engagement workflows. Achieved 99.8% system uptime and 70% performance improvement through strategic caching and query optimization.',
            'tech': 'Java, Spring Boot, MySQL, AWS S3, ELK Stack, Redis, REST APIs'
        },
        {
            'title': 'NEXOGIC - HEALTH PRACTITIONER NETWORK | Oodles Technologies',
            'desc': 'Developed HIPAA-compliant healthcare collaboration platform enabling secure practitioner networking with encrypted chat, WebRTC video conferencing, and MFA authentication. Built research sharing infrastructure serving 10K+ healthcare professionals with zero security breaches.',
            'tech': 'Java, Spring Boot, PostgreSQL, AWS S3, WebRTC, OAuth 2.0, JWT'
        },
        {
            'title': 'BELFRICS & PARITEX CRYPTOCURRENCY EXCHANGE | Oodles Technologies',
            'desc': 'Engineered high-frequency cryptocurrency trading platform handling real-time order matching and settlement for Bitcoin, Ethereum, and digital assets. Implemented fault-tolerant architecture with Redis caching achieving sub-100ms trade execution latency under peak load.',
            'tech': 'Java, Spring Boot, MySQL, Redis, AWS, Blockchain, WebSocket, REST APIs'
        }
    ]
    
    for project in projects:
        p = doc.add_paragraph()
        run = p.add_run(project['title'])
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 82, 130)
        
        p = doc.add_paragraph(project['desc'])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        p = doc.add_paragraph()
        run = p.add_run(f"Technologies: {project['tech']}")
        run.italic = True
    
    # Education
    heading = doc.add_heading('EDUCATION & CERTIFICATIONS', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    education = [
        ('Master of Computer Applications (MCA)', 'Aug 2016 - Jul 2019', 'M.J.P. Rohilkhand University, Bareilly, India'),
        ('Bachelor of Science (B.Sc.)', 'Jul 2010 - Dec 2013', 'M.J.P. Rohilkhand University, Bareilly, India'),
        ('AWS Certified Cloud Practitioner (CFL-C01)', '2023', 'Amazon Web Services')
    ]
    
    for degree, date, institution in education:
        p = doc.add_paragraph()
        run = p.add_run(f'{degree} | {date}')
        run.font.bold = True
        p = doc.add_paragraph(institution)
    
    # Key Achievements
    heading = doc.add_heading('KEY ACHIEVEMENTS & METRICS', level=1)
    heading.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    achievements = [
        'Delivered $500K+ in cumulative cost savings through cloud optimization and infrastructure modernization',
        'Architected 15+ microservices processing 1M+ daily transactions with 99.95% uptime SLA',
        'Achieved 80% response time improvement (1.2s to 240ms) through strategic performance optimization',
        'Reduced infrastructure costs by $200K annually (40%) through serverless architecture adoption',
        'Led teams of 15+ developers across 3 global locations with 98% on-time delivery rate',
        'Improved test coverage from 45% to 92% and code maintainability from 40% to 85%',
        'Reduced production incidents by 70% and MTTR from 4 hours to 45 minutes',
        'Conducted 500+ code reviews and mentored junior developers on best practices'
    ]
    
    for achievement in achievements:
        doc.add_paragraph(f'✓ {achievement}', style='List Bullet')
    
    # Save document
    filename = 'Lokesh_Babu_Sharma_Resume.docx'
    doc.save(filename)
    print(f"DOCX resume generated successfully: {filename}")

if __name__ == "__main__":
    create_resume_docx()
